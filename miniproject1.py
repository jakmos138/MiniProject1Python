import warnings
import re
import ebooklib
from ebooklib import epub
import requests
from PIL import Image
from io import BytesIO
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Cm
from collections import Counter


warnings.filterwarnings("ignore", message="In the future version we will turn default option ignore_ncx to True.")
warnings.filterwarnings("ignore", message="This search incorrectly ignores the root element, and will be fixed in a future version.")


def read_epub_file(file_path):
    book = epub.read_epub(file_path)
    return book


def get_book_info(book):
    title = ""
    authors = []

    for metadata in book.get_metadata('DC', 'title'):
        title = f'{metadata[0]}'

    for metadata in book.get_metadata('DC', 'creator'):
        authors.append(f'{metadata[0]}')

    return title, authors


def get_first_chapter(book, start_pattern, end_pattern):
    first_chapter_content = None
    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            content = item.get_body_content()
            break

    text = content.decode('utf-8') if content else None

    pattern = re.compile(
        rf'{re.escape(start_pattern)}(.*?)\s*{re.escape(end_pattern)}',
        re.DOTALL)
    match = pattern.search(text)
    if match:
        return match.group(1)
    else:
        return None


def decode_paragraphs(text):
    paragraphs = text.replace("\n", " ").split('</p>')
    clean_paragraphs = [re.sub(r'<[^>]*>', '', paragraph).strip()
                        for paragraph in paragraphs
                        if paragraph.strip()]
    return clean_paragraphs[1:]


def evaluate_paragraph_lengths(paragraphs):
    paragraph_lengths = [len(paragraph.split()) for paragraph in paragraphs]
    return paragraph_lengths


def plot_paragraph_lengths(paragraph_lengths):
    counter = Counter(paragraph_lengths)
    lengths = list(counter.keys())
    frequencies = list(counter.values())

    plt.figure(figsize=(10, 6))
    plt.bar(lengths, frequencies, width=1.0, edgecolor='black')
    plt.xlabel('Number of Words in Paragraph')
    plt.ylabel('Number of Paragraphs')
    plt.title('Distribution of Paragraph Lengths in First Chapter')
    plt.savefig('paragraph_lengths_distribution.png')
    plt.close()


def download_image(url):
    try:
        response = requests.get(url)
        response.raise_for_status()  # Check if the request was successful
        return Image.open(BytesIO(response.content))
    except requests.exceptions.RequestException as e:
        print(f"Error downloading the image: {e}")
        return None
    except IOError as e:
        print(f"Error opening the image: {e}")
        return None


def combine_images(img_url, img_path):
    img1 = download_image(img_url)
    if img1 is None:
        raise SystemExit("Failed to download Picture #1")

    original_width, original_height = img1.size
    new_width, new_height = 1000, 700

    left = (original_width - new_width) / 2
    top = (original_height - new_height) / 2
    right = (original_width + new_width) / 2
    bottom = (original_height + new_height) / 2

    crop_box = (left, top, right, bottom)
    img1_cropped = img1.crop(crop_box)

    resize_dimensions = (800, 700)
    img1_resized = img1_cropped.resize(resize_dimensions)

    img2 = Image.open(img_path)

    # Rotate the image by a chosen angle
    angle = 45  # Example value
    img2_rotated = img2.rotate(angle, expand=True)

    offset_x = 150  # Example offset value
    offset_y = 150  # Example offset value
    position = ((img1_cropped.width - img2_rotated.width) // 2 + offset_x,
                (img1_cropped.height - img2_rotated.height) // 2 + offset_y)

    # Create a new image to avoid modifying the original
    final_image = img1_resized.copy()
    final_image.paste(img2_rotated, position, img2_rotated)

    # Save the final image
    final_image.save("final_image.png")

    return final_image


def create_document(title, img_path, authors, report_author, plot_path, description):
    doc = Document()
    
    # Title page
    doc.add_heading(title, level=1)
    doc.add_picture(img_path, width=Cm(12))
    doc.add_paragraph(f"Author: {', '.join(authors)}")
    doc.add_paragraph(f"Report Author: {report_author}")
    
    doc.add_page_break()
    
    # Info page
    doc.add_heading('Paragraph Lengths Distribution', level=1)
    doc.add_picture(plot_path, width=Cm(12))
    
    doc.add_paragraph(description)
    
    doc.save('report.docx')


def main():
    file_path = 'book.epub'  # Path to your ePub file
    start_pattern = '<p class="ph1">'
    end_pattern = '<hr class="chap"/>'
    img_url = "https://cdn.thecollector.com/wp-content/uploads/2023/04/hp-lovecraft-cthulhu-mythos.jpg"
    img_path = "img2.png"
    project_author = "Jakub Moszyński"
    # Read the ePub file
    book = read_epub_file(file_path)

    # Display the structure of the book
    title, authors = get_book_info(book)

    first_chapter_text = get_first_chapter(book, start_pattern, end_pattern)

    first_chapter_paragraphs = decode_paragraphs(first_chapter_text)

    # Evaluate paragraph lengths
    paragraph_lengths = evaluate_paragraph_lengths(first_chapter_paragraphs)

    # Create plot
    plot_paragraph_lengths(paragraph_lengths)

    final_image = combine_images(img_url, img_path)

    description = (
        f"The first chapter contains {len(first_chapter_paragraphs)} paragraphs.\n"
        f"The total number of words in the first chapter is {sum(paragraph_lengths)}.\n"
        f"The minimal number of words in a paragraph is {min(paragraph_lengths)}.\n"
        f"The maximal number of words in a paragraph is {max(paragraph_lengths)}.\n"
        f"The average number of words per paragraph is {sum(paragraph_lengths)/len(paragraph_lengths):.2f}."
    )

    create_document(title, "final_image.png", authors, project_author, "paragraph_lengths_distribution.png", description)


if __name__ == '__main__':
    main()
